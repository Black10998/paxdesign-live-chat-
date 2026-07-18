#!/usr/bin/env python3
"""Build professional DOCX/PDF for Guideline 3.2 attachment (content unchanged)."""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs/app-store/GUIDELINE-3.2-ATTACHMENT-BUILD-159.md"
EXPORT_DIR = ROOT / "docs/app-store/exports"
DOCX_PATH = EXPORT_DIR / "PAXDesign-Guideline-3.2-Attachment-Build-159.docx"
PDF_PATH = EXPORT_DIR / "PAXDesign-Guideline-3.2-Attachment-Build-159.pdf"


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run2 = paragraph.add_run("1")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run2 = paragraph.add_run("Right-click and choose Update Field if the table of contents is empty.")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_end)


def set_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    add_page_number(p)


def add_formatted_run(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = [line.strip().strip("|") for line in lines if line.strip()]
    header = [cell.strip() for cell in rows[0].split("|")]
    body = []
    for row in rows[2:]:
        body.append([cell.strip() for cell in row.split("|")])
    return header, body


def build_docx() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if level == 1:
            hs.font.size = Pt(22)
        elif level == 2:
            hs.font.size = Pt(16)
        else:
            hs.font.size = Pt(13)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = "Calibri"
            i += 1
            continue

        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading == "Purpose of This Document":
                doc.add_page_break()
                doc.add_heading("Table of Contents", level=1)
                toc_p = doc.add_paragraph()
                add_toc(toc_p)
                doc.add_page_break()
            doc.add_heading(heading, level=1 if re.match(r"^\d+\.", heading) else 2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, body = parse_table(table_lines)
            table = doc.add_table(rows=1 + len(body), cols=len(header))
            table.style = "Table Grid"
            for col, value in enumerate(header):
                cell = table.rows[0].cells[col]
                cell.text = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for r, row in enumerate(body, start=1):
                for c, value in enumerate(row):
                    table.rows[r].cells[c].text = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
            doc.add_paragraph()
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and ":**" in stripped:
            p = doc.add_paragraph()
            add_formatted_run(p, stripped)
            i += 1
            continue

        if stripped:
            p = doc.add_paragraph()
            add_formatted_run(p, stripped)
        i += 1

    set_footer(doc)
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def build_pdf() -> None:
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(EXPORT_DIR),
            str(DOCX_PATH),
        ],
        check=True,
        capture_output=True,
        text=True,
    )


def main() -> int:
    build_docx()
    build_pdf()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
